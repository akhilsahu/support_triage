"""
Document Parser — universal text extractor for the RAG pipeline.

Supported formats:
  .pdf    — pypdf (text-based PDFs)
  .docx   — python-docx
  .doc    — python-docx (best-effort; .doc may need libreoffice)
  .txt    — plain text (any encoding)
  .md     — Markdown (strips fences, returns raw text)
  .html   — BeautifulSoup tag stripping
  .htm    — same as .html
  .json   — pretty-printed key/value text
  .csv    — row-by-row text

Each parser returns a list of ParsedPage objects so chunking always
knows which page/section content came from.
"""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import structlog

logger = structlog.get_logger()


@dataclass
class ParsedPage:
    """A logical page or section extracted from a document."""
    page: int          # 1-based
    text: str
    section: str = ""  # heading / sheet name / JSON key path etc.


@dataclass
class ParsedDocument:
    filename: str
    extension: str
    pages: List[ParsedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def page_count(self) -> int:
        return len(self.pages)


# ── Public entry point ────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".txt", ".md", ".rst",
    ".html", ".htm",
    ".json", ".jsonl",
    ".csv",
}


def parse(raw: bytes, filename: str) -> ParsedDocument:
    """
    Parse raw file bytes into a ParsedDocument.

    Args:
        raw:      File bytes.
        filename: Original filename (used to detect format).

    Returns:
        ParsedDocument with per-page text and metadata.

    Raises:
        ValueError: Unsupported file extension.
        RuntimeError: Required library missing or extraction failed.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(raw, filename)
    if ext in (".docx", ".doc"):
        return _parse_docx(raw, filename)
    if ext in (".txt", ".md", ".rst"):
        return _parse_text(raw, filename)
    if ext in (".html", ".htm"):
        return _parse_html(raw, filename)
    if ext in (".json", ".jsonl"):
        return _parse_json(raw, filename)
    if ext == ".csv":
        return _parse_csv(raw, filename)

    raise ValueError(
        f"Unsupported file type '{ext}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(raw: bytes, filename: str) -> ParsedDocument:
    """
    Parse PDF using PyMuPDF. Font-size analysis detects headings so each
    logical section becomes its own ParsedPage with section= set.
    Falls back to pypdf if PyMuPDF is not installed.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=raw, filetype="pdf")
        meta = {
            "title":      doc.metadata.get("title"),
            "author":     doc.metadata.get("author"),
            "page_count": doc.page_count,
        }

        # Collect all spans across all pages to determine body font size
        all_spans: list[dict] = []
        for pi in range(doc.page_count):
            for block in doc.load_page(pi).get_text("dict")["blocks"]:
                if block["type"] != 0:  # 0 = text block
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["text"].strip():
                            all_spans.append(span)

        # Body font = most common font size (mode)
        if all_spans:
            from collections import Counter
            size_counts = Counter(round(s["size"]) for s in all_spans)
            body_size = size_counts.most_common(1)[0][0]
        else:
            body_size = 11

        heading_threshold = body_size * 1.15  # 15% larger than body = heading

        # Build ParsedPage list: new section on each heading-like span
        pages: list[ParsedPage] = []
        current_section = ""
        buffer: list[str] = []
        logical_page = 1

        def _flush():
            nonlocal logical_page
            text = "\n".join(buffer).strip()
            if text:
                pages.append(ParsedPage(page=logical_page, text=text, section=current_section))
                logical_page += 1
            buffer.clear()

        for page_idx in range(doc.page_count):
            for block in doc.load_page(page_idx).get_text("dict")["blocks"]:
                if block["type"] != 0:
                    continue
                for line in block["lines"]:
                    line_text = " ".join(s["text"] for s in line["spans"]).strip()
                    if not line_text:
                        continue
                    # Detect heading: larger font or bold, short line, no terminal punct
                    span = line["spans"][0] if line["spans"] else {}
                    is_heading = (
                        round(span.get("size", 0)) >= heading_threshold
                        and len(line_text.split()) <= 15
                        and not line_text[-1] in ".!?,:;"
                    )
                    if is_heading:
                        _flush()
                        current_section = line_text
                    else:
                        buffer.append(line_text)

        _flush()

        # If no headings were found, fall back to one ParsedPage per PDF page
        if not pages:
            pages = []
            for i in range(doc.page_count):
                text = doc.load_page(i).get_text().strip()
                if text:
                    pages.append(ParsedPage(page=i + 1, text=text))

        logger.info("PDF parsed", filename=filename, pages=len(pages), body_font=body_size)
        return ParsedDocument(filename=filename, extension=".pdf", pages=pages, metadata=meta)

    except ImportError:
        pass

    # Fallback: pypdf (no heading detection)
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(ParsedPage(page=i, text=text.strip()))
        meta = {
            "author":     reader.metadata.author if reader.metadata else None,
            "title":      reader.metadata.title if reader.metadata else None,
            "page_count": len(reader.pages),
        }
        logger.info("PDF parsed (pypdf fallback)", filename=filename, pages=len(pages))
        return ParsedDocument(filename=filename, extension=".pdf", pages=pages, metadata=meta)
    except ImportError:
        raise RuntimeError("pymupdf or pypdf not installed. Run: pip install pymupdf")


# ── DOCX / DOC ────────────────────────────────────────────────────────────────

def _parse_docx(raw: bytes, filename: str) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(raw))

        pages: List[ParsedPage] = []
        current_section = ""
        buffer: List[str] = []
        page = 1

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            # Headings start a new logical page/section
            if style.startswith("Heading"):
                if buffer:
                    pages.append(ParsedPage(page=page, text="\n".join(buffer), section=current_section))
                    page += 1
                    buffer = []
                current_section = text
            else:
                buffer.append(text)

        if buffer:
            pages.append(ParsedPage(page=page, text="\n".join(buffer), section=current_section))

        # If no headings at all, treat whole doc as one page
        if not pages and doc.paragraphs:
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            pages = [ParsedPage(page=1, text=all_text)]

        # Also extract tables
        for tbl_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append(" | ".join(c.text.strip() for c in row.cells))
            if rows:
                pages.append(ParsedPage(
                    page=page + tbl_idx + 1,
                    text="\n".join(rows),
                    section=f"Table {tbl_idx + 1}"
                ))

        meta = {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        logger.info("DOCX parsed", filename=filename, sections=len(pages))
        return ParsedDocument(filename=filename, extension=Path(filename).suffix.lower(), pages=pages, metadata=meta)

    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


# ── Plain text / Markdown ─────────────────────────────────────────────────────

def _parse_text(raw: bytes, filename: str) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    text = _decode(raw)

    if ext == ".md":
        return _parse_markdown(text, filename)

    # Plain .txt / .rst — split on double newlines, group ~400 words per page
    paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    pages: List[ParsedPage] = []
    buffer: List[str] = []
    word_count = 0
    page = 1

    for para in paras:
        buffer.append(para)
        word_count += len(para.split())
        if word_count >= 400:
            pages.append(ParsedPage(page=page, text="\n\n".join(buffer)))
            page += 1
            buffer = []
            word_count = 0

    if buffer:
        pages.append(ParsedPage(page=page, text="\n\n".join(buffer)))

    if not pages:
        pages = [ParsedPage(page=1, text=text.strip())]

    logger.info("Text parsed", filename=filename, pages=len(pages))
    return ParsedDocument(filename=filename, extension=ext, pages=pages, metadata={"char_count": len(text)})


def _parse_markdown(text: str, filename: str) -> ParsedDocument:
    """
    Parse Markdown using markdown-it-py token stream.
    Each heading starts a new ParsedPage with section= set to the heading text.
    Body text accumulates under the current heading.
    """
    try:
        from markdown_it import MarkdownIt
        md = MarkdownIt()
        tokens = md.parse(text)

        pages: List[ParsedPage] = []
        current_section = ""
        buffer: List[str] = []
        page = 1

        def _flush():
            nonlocal page
            body = "\n\n".join(buffer).strip()
            if body:
                pages.append(ParsedPage(page=page, text=body, section=current_section))
                page += 1
            buffer.clear()

        i = 0
        while i < len(tokens):
            tok = tokens[i]
            if tok.type == "heading_open":
                _flush()
                # Next token is inline with the heading text
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    current_section = tokens[i + 1].content.strip()
                    i += 2  # skip heading_open + inline
                else:
                    i += 1
            elif tok.type == "inline":
                if tok.content.strip():
                    buffer.append(tok.content.strip())
                i += 1
            elif tok.type == "fence":
                # Code block — include as-is (strip the fence markers)
                buffer.append(tok.content.strip())
                i += 1
            else:
                i += 1

        _flush()

        if not pages:
            pages = [ParsedPage(page=1, text=text.strip())]

        logger.info("Markdown parsed", filename=filename, sections=len(pages))
        return ParsedDocument(filename=filename, extension=".md", pages=pages, metadata={"char_count": len(text)})

    except ImportError:
        # Fallback: regex-based heading split
        parts = re.split(r"(?m)^(#{1,3} .+)$", text)
        pages = []
        current_section = ""
        page = 1
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if re.match(r"^#{1,3} ", part):
                current_section = re.sub(r"^#+\s*", "", part)
            elif part:
                pages.append(ParsedPage(page=page, text=part, section=current_section))
                page += 1
        if not pages:
            pages = [ParsedPage(page=1, text=text.strip())]
        logger.info("Markdown parsed (fallback)", filename=filename, sections=len(pages))
        return ParsedDocument(filename=filename, extension=".md", pages=pages)


# ── HTML ──────────────────────────────────────────────────────────────────────

def _parse_html(raw: bytes, filename: str) -> ParsedDocument:
    """
    Parse HTML using BeautifulSoup. Each h1/h2/h3 tag starts a new ParsedPage
    with section= set to the heading text. Body text accumulates under the
    current heading. Falls back to flat text split if bs4 is unavailable.
    """
    try:
        from bs4 import BeautifulSoup, Tag

        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        pages: List[ParsedPage] = []
        current_section = ""
        buffer: List[str] = []
        page = 1

        def _flush():
            nonlocal page
            body = "\n\n".join(buffer).strip()
            if body:
                pages.append(ParsedPage(page=page, text=body, section=current_section))
                page += 1
            buffer.clear()

        body_tag = soup.find("body") or soup
        for element in body_tag.descendants:  # type: ignore[union-attr]
            if not isinstance(element, Tag):
                continue
            if element.name in ("h1", "h2", "h3"):
                _flush()
                current_section = element.get_text(separator=" ").strip()
            elif element.name in ("p", "li", "td", "th", "blockquote", "pre"):
                text = element.get_text(separator=" ").strip()
                if text:
                    buffer.append(text)

        _flush()

        if not pages:
            flat = soup.get_text(separator="\n")
            pages = [ParsedPage(page=1, text=flat.strip())]

        logger.info("HTML parsed", filename=filename, sections=len(pages))
        return ParsedDocument(filename=filename, extension=".html", pages=pages)

    except ImportError:
        pass

    # Regex fallback (no bs4)
    text = _decode(raw)
    text = re.sub(r"<script[\s\S]*?</script>", "", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>",  "", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+", " ", text).strip()
    pages = [ParsedPage(page=1, text=text)]
    logger.info("HTML parsed (regex fallback)", filename=filename)
    return ParsedDocument(filename=filename, extension=".html", pages=pages)


# ── JSON ──────────────────────────────────────────────────────────────────────

def _parse_json(raw: bytes, filename: str) -> ParsedDocument:
    text = _decode(raw)
    ext = Path(filename).suffix.lower()
    pages: List[ParsedPage] = []

    if ext == ".jsonl":
        for i, line in enumerate(text.splitlines(), 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                pages.append(ParsedPage(page=i, text=_flatten_json(obj), section=f"Record {i}"))
            except json.JSONDecodeError:
                pages.append(ParsedPage(page=i, text=line))
    else:
        try:
            data = json.loads(text)
        except json.JSONDecodeError as e:
            raise RuntimeError(f"Invalid JSON: {e}")

        if isinstance(data, list):
            for i, item in enumerate(data, 1):
                pages.append(ParsedPage(page=i, text=_flatten_json(item), section=f"Item {i}"))
        else:
            # Split top-level keys into logical sections
            for i, (key, val) in enumerate(data.items(), 1):
                pages.append(ParsedPage(page=i, text=_flatten_json({key: val}), section=key))

    if not pages:
        pages = [ParsedPage(page=1, text=text[:5000])]

    logger.info("JSON parsed", filename=filename, records=len(pages))
    return ParsedDocument(filename=filename, extension=ext, pages=pages)


def _flatten_json(obj, prefix: str = "", sep: str = " → ") -> str:
    lines = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}{sep}{k}" if prefix else k
            if isinstance(v, (dict, list)):
                lines.append(_flatten_json(v, key, sep))
            else:
                lines.append(f"{key}: {v}")
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            lines.append(_flatten_json(item, f"{prefix}[{i}]", sep))
    else:
        lines.append(str(obj))
    return "\n".join(lines)


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(raw: bytes, filename: str) -> ParsedDocument:
    text = _decode(raw)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ParsedDocument(filename=filename, extension=".csv", pages=[])

    headers = rows[0]
    pages: List[ParsedPage] = []
    chunk_rows: List[str] = [" | ".join(headers)]
    page = 1

    for row_idx, row in enumerate(rows[1:], 1):
        chunk_rows.append(" | ".join(row))
        if row_idx % 50 == 0:  # 50 rows per page
            pages.append(ParsedPage(page=page, text="\n".join(chunk_rows), section=f"Rows {page*50-49}–{page*50}"))
            page += 1
            chunk_rows = [" | ".join(headers)]

    if len(chunk_rows) > 1:
        pages.append(ParsedPage(page=page, text="\n".join(chunk_rows)))

    logger.info("CSV parsed", filename=filename, rows=len(rows) - 1, pages=len(pages))
    return ParsedDocument(
        filename=filename, extension=".csv", pages=pages,
        metadata={"row_count": len(rows) - 1, "columns": headers}
    )


# ── Chunk dataclass (used by both document_parser and chunking modules) ────────

@dataclass
class Chunk:
    text: str
    page: int
    chunk_index: int
    section: str = ""


def chunk_document(doc: ParsedDocument, _chunk_size: int = 1000, _overlap: int = 150) -> List[Chunk]:
    """
    Convenience wrapper — delegates to chunking.chunk() with the
    extension-derived config.  Kept for backwards compatibility.
    """
    from app.rag.chunking import chunk as _chunk
    return _chunk(doc)


# ── Utility ───────────────────────────────────────────────────────────────────

def _decode(raw: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")
