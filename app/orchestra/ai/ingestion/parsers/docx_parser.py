"""
DocxParser — python-docx with native table extraction.
Falls back to LibreOffice conversion for legacy .doc binary format.
"""

from __future__ import annotations

import io
import subprocess
import tempfile
from pathlib import Path
from typing import List

import structlog

from app.rag.document_parser import ParsedDocument, ParsedPage
from app.orchestra.ai.ingestion.core.base import BaseParser

logger = structlog.get_logger()


class DocxParser(BaseParser):
    extensions = [".docx", ".doc"]

    def parse(self, raw: bytes, filename: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        if ext == ".doc":
            raw, filename = self._convert_doc(raw, filename)

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc     = DocxDocument(io.BytesIO(raw))
        pages: List[ParsedPage] = []
        section = ""
        buffer: List[str] = []
        page    = 1

        def _flush():
            nonlocal page
            text = "\n".join(buffer).strip()
            if text:
                pages.append(ParsedPage(page=page, text=text, section=section))
                page += 1
            buffer.clear()

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text  = para.text.strip()
            if not text:
                continue
            if style.startswith("Heading"):
                _flush()
                section = text
            else:
                buffer.append(text)

        _flush()

        # Each table → its own ParsedPage with pipe-delimited rows
        for tbl_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                # Deduplicate merged cells (python-docx repeats them)
                deduped = []
                for i, cell in enumerate(cells):
                    if i == 0 or cell != cells[i - 1]:
                        deduped.append(cell)
                rows.append(" | ".join(deduped))

            if rows:
                label = f"Table {tbl_idx + 1}" + (f" ({section})" if section else "")
                pages.append(ParsedPage(
                    page=page + tbl_idx,
                    text=label + ":\n" + "\n".join(rows),
                    section=label,
                    is_table=True,
                ))

        if not pages:
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if all_text:
                pages = [ParsedPage(page=1, text=all_text)]

        meta = {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        logger.info("ingestion.docx", filename=filename, sections=len(pages), tables=len(doc.tables))
        return ParsedDocument(filename=filename, extension=Path(filename).suffix.lower(),
                              pages=pages, metadata=meta)

    def _convert_doc(self, raw: bytes, filename: str):
        """Convert .doc → .docx via LibreOffice headless."""
        if not self.cfg.libreoffice_enabled:
            raise ValueError(".doc files require LibreOffice. Set INGESTION_LIBREOFFICE_ENABLED=true")

        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / filename
            src.write_bytes(raw)
            try:
                subprocess.run(
                    [self.cfg.libreoffice_path, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, str(src)],
                    check=True, capture_output=True, timeout=60,
                )
            except (subprocess.CalledProcessError, FileNotFoundError) as e:
                raise RuntimeError(
                    f"LibreOffice conversion failed for '{filename}'. "
                    f"Install LibreOffice or set INGESTION_LIBREOFFICE_PATH. Error: {e}"
                )
            out = src.with_suffix(".docx")
            if not out.exists():
                raise RuntimeError(f"LibreOffice did not produce output for '{filename}'")
            logger.info("ingestion.doc.converted", filename=filename)
            return out.read_bytes(), out.name
